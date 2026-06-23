from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
import time
from PIL import Image
import io
import cv2
import numpy as np
import os
from docx import Document
from docx.shared import Inches
import requests
from io import BytesIO
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement


# -------------------- 浏览器设置 --------------------
def create_browser():
    options = Options()
    # 创建chrome浏览器的配置对象
    options.add_argument('--no-sandbox')
    options.add_experimental_option('detach', True)
    # 程序结束之后 true 不会关闭， false 会自动关闭
    return webdriver.Chrome(options=options)

# -------------------- 绘制多边形并保存图片 --------------------
def draw_polygon_on_image(pil_image):
    """
    在 PIL.Image 上用鼠标手动绘制多边形
    返回：
        image_with_boundary (PIL.Image)
        points (list of (x, y))
    """
    # PIL → OpenCV
    image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)

    points = []
    drawing = True

    def mouse_callback(event, x, y, flags, param):
        nonlocal points, drawing
        if event == cv2.EVENT_LBUTTONDOWN:
            points.append((x, y))
        elif event == cv2.EVENT_RBUTTONDOWN:
            drawing = False

    cv2.namedWindow("Draw Boundary")
    cv2.setMouseCallback("Draw Boundary", mouse_callback)

    while drawing:
        temp = image.copy()
        if len(points) > 1:
            cv2.polylines(
                temp,
                [np.array(points, dtype=np.int32)],
                False,
                (0, 0, 255),
                2
            )
        cv2.imshow("Draw Boundary", temp)

        # ESC 也可以退出
        if cv2.waitKey(1) & 0xFF == 27:
            drawing = False

    cv2.destroyAllWindows()

    if len(points) < 3:
        raise ValueError("At least 3 points required to define polygon")

    # 最终画到原图
    cv2.polylines(
        image,
        [np.array(points, dtype=np.int32)],
        False,
        (0, 0, 255),
        2
    )

    # OpenCV → PIL
    image_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
    image_pil_out = Image.fromarray(image_rgb)

    return image_pil_out, points




# -------------------- 检测颜色存在 --------------------
def color_in_region(image_rgb, points, target_color, tol=60):
    mask = np.zeros(image_rgb.shape[:2], dtype=np.uint8)
    cv2.fillPoly(mask, [np.array(points, dtype=np.int32)], 255)
    masked_pixels = image_rgb[mask==255]
    diff = np.abs(masked_pixels - target_color)
    match = np.all(diff <= tol, axis=1)
    return np.any(match)

# -------------------- 解析 Postcode 获取坐标 --------------------
def postcode_to_easting_northing(postcode: str):
    postcode = postcode.strip().replace(" ", "")
    url = f"https://api.postcodes.io/postcodes/{postcode}"
    response = requests.get(url)
    if response.status_code != 200:
        return None, None
    data = response.json()['result']
    return data['eastings'], data['northings']

def insert_paragraph_and_image(doc, placeholder, text, image_file, img_width=4):
    target_paragraph = None
    for p in doc.paragraphs:
        if placeholder in p.text:
            target_paragraph = p
            break
    if target_paragraph is None:
        raise ValueError(f"未找到占位符 {placeholder}")

    # 插入文字
    text_p = doc.add_paragraph(text)
    text_p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    target_paragraph._element.addnext(text_p._element)

    # 插入图片
    if os.path.exists(image_file):
        pic_p = doc.add_paragraph()
        pic_run = pic_p.add_run()
        pic_run.add_picture(image_file, width=Inches(img_width))
        pic_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        text_p._element.addnext(pic_p._element)

    # 删除占位符段落
    target_paragraph._element.getparent().remove(target_paragraph._element)

def overlay_image_top_right(
    base_image_path,
    overlay_image_path,
    output_path,
    scale=0.4,      # 占底图宽度的比例（0.2 ≈ 1/5）
    margin=20,      # 距离边缘的像素
    alpha=0.9       # 透明度
):
    base = cv2.imread(base_image_path, cv2.IMREAD_UNCHANGED)
    overlay = cv2.imread(overlay_image_path, cv2.IMREAD_UNCHANGED)

    if base is None or overlay is None:
        raise ValueError("Base image or overlay image not found")

    h, w = base.shape[:2]

    # 计算 overlay 尺寸（按底图宽度的 1/5 缩放）
    new_w = int(w * scale)
    ratio = new_w / overlay.shape[1]
    new_h = int(overlay.shape[0] * ratio)

    overlay = cv2.resize(overlay, (new_w, new_h))

    # 右上角位置
    x = w - new_w - margin
    y = margin

    # 如果有 alpha 通道 透明
    if overlay.shape[2] == 4:
        overlay_rgb = overlay[:, :, :3]
        overlay_alpha = (overlay[:, :, 3] / 255.0) * alpha

        for c in range(3):
            base[y:y+new_h, x:x+new_w, c] = (
                overlay_alpha * overlay_rgb[:, :, c] +
                (1 - overlay_alpha) * base[y:y+new_h, x:x+new_w, c]
            )
    else:
        roi = base[y:y+new_h, x:x+new_w]
        blended = cv2.addWeighted(overlay, alpha, roi, 1 - alpha, 0)
        base[y:y+new_h, x:x+new_w] = blended

    cv2.imwrite(output_path, base)


# -------------------- 信息输入 --------------------
def replace_placeholder_all(doc, placeholder, text):
    # 1. 替换段落
    for p in doc.paragraphs:
        for r in p.runs:
            if placeholder in r.text:
                r.text = r.text.replace(placeholder, text)

    # 2. 替换表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholder_all(cell, placeholder, text)  # 递归

    # 3. 替换页眉页脚
    if hasattr(doc, "sections"):
        for section in doc.sections:
            header = section.header
            footer = section.footer
            replace_placeholder_all(header, placeholder, text)
            replace_placeholder_all(footer, placeholder, text)



# -------------------- 插入 --------------------  (关于带图片的插入不对)
def insert_text_and_image_below(doc, placeholder=None, text=None, image=None, img_width=4):
    # 找到占位符段落
    target_paragraph = None
    if placeholder:
        for p in doc.paragraphs:
            if placeholder in p.text:
                target_paragraph = p
                break
        if target_paragraph is None:
            raise ValueError(f"未找到占位符 {placeholder}")
    else:
        target_paragraph = doc.paragraphs[-1]

    last_p = target_paragraph  # 用于返回最后插入的段落

    # 插入文字
    if text:
        # 在 target_paragraph 后插入文字段落
        new_p_elm = OxmlElement('w:p')
        target_paragraph._element.addnext(new_p_elm)
        text_p = doc.add_paragraph()
        text_p._element = new_p_elm
        text_p.add_run(text)
        last_p = text_p  # 更新最后段落

    # 插入图片
    if image:
        if image.mode != 'RGB':
            image = image.convert('RGB')
        img_byte_arr = BytesIO()
        image.save(img_byte_arr, format='PNG')
        img_byte_arr.seek(0)

        # 在 last_p 后插入图片段落
        new_p_elm = OxmlElement('w:p')
        last_p._element.addnext(new_p_elm)
        pic_p = doc.add_paragraph()
        pic_p._element = new_p_elm
        run = pic_p.add_run()
        run.add_picture(img_byte_arr, width=Inches(img_width))
        pic_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        last_p = pic_p  # 更新最后段落

    # 删除占位符段落（如果存在 placeholder）
    if placeholder and target_paragraph in doc.paragraphs:
        target_paragraph._element.getparent().remove(target_paragraph._element)

    return last_p



# -------------------- 叠加图片--------------------
def overlay_image_top_right_pil(
    base_pil,
    overlay_pil,
    scale=0.4,
    margin=20,
    alpha=0.95
):
    base = cv2.cvtColor(np.array(base_pil), cv2.COLOR_RGB2BGR)
    overlay = cv2.cvtColor(np.array(overlay_pil), cv2.COLOR_RGB2BGR)

    h, w = base.shape[:2]

    new_w = int(w * scale)
    ratio = new_w / overlay.shape[1]
    new_h = int(overlay.shape[0] * ratio)
    overlay = cv2.resize(overlay, (new_w, new_h))

    x = w - new_w - margin
    y = margin

    blended = cv2.addWeighted(
        overlay,
        alpha,
        base[y:y+new_h, x:x+new_w],
        1 - alpha,
        0
    )
    base[y:y+new_h, x:x+new_w] = blended

    return Image.fromarray(cv2.cvtColor(base, cv2.COLOR_BGR2RGB))




# -------------------- 主程序 --------------------

# ---------- 用户输入项目信息 ----------
project_name = input("Project Name: ")
client_name = input("Client Name: ")
project_location = input("Project Location: ")
project_number = input("Project Number: ")


if __name__ == "__main__":
    postcode = input("Postcode: ")

    # ---------- 1. Fluvial/Tidal Flood Analysis ----------
    browser = create_browser()
    browser.set_window_position(-2000,0)
    browser.get('https://flood-map-for-planning.service.gov.uk/map')
    browser.set_window_size(1920,1200)
    time.sleep(2)

    # 搜索邮编
    browser.find_element(By.XPATH, '//*[@id="map"]/div/div/div[2]/div[3]/div[1]/button[1]').click()
    search_box = browser.find_element(By.ID, 'map-search')
    search_box.send_keys(postcode)
    search_box.send_keys(Keys.ENTER)
    time.sleep(3)
    browser.find_element(By.XPATH, '/html/body/div/div[1]/div[2]/button[1]').click()
    browser.find_element(By.XPATH, '/html/body/div/div[2]/div[2]/button').click()
    time.sleep(5)
    # 放大
    browser.find_element(By.XPATH, '//*[@id="map"]/div/div[2]/div[2]/div[3]/div[4]/div[3]/div[1]/button').click()
    time.sleep(2)
    # 叉掉legend
    browser.find_element(By.XPATH, '//*[@id="map-panel-key"]/div[1]/button').click()
    time.sleep(1)

    map_el = browser.find_element(By.ID, "map")
    png = browser.get_screenshot_as_png()

    image = Image.open(io.BytesIO(png))
    loc, size = map_el.location, map_el.size

    left, top = loc['x']+600, loc['y']+100
    right, bottom = loc['x']+size['width'], loc['y']+size['height']

    fluvial_crop = image.crop((left, top, right, bottom))

    browser.quit()
    # ---------- 叠加 Flood Zone 图例 ----------
    legend_img = Image.open("Floodzone.png")
    fluvial_with_legend = overlay_image_top_right_pil(
        fluvial_crop,
        legend_img,
        scale=0.4,
        alpha=0.95
    )

    # ---------- 手动画边界 ----------

    fluvial_img_pil, fluvial_points = draw_polygon_on_image(fluvial_with_legend)

    fluvial_rgb = np.array(fluvial_img_pil)

    flood2_color = np.array([29,112,184])
    flood3_color = np.array([0,48,120])
    if color_in_region(fluvial_rgb, fluvial_points, flood3_color):
        flood_zone = "Flood Zone 3"
    elif color_in_region(fluvial_rgb, fluvial_points, flood2_color):
        flood_zone = "Flood Zone 2"
    else:
        flood_zone = "Flood Zone 1"

    risk_dict = {
        "Flood Zone 1": ("low risk", "low"),
        "Flood Zone 2": ("moderate risk", "medium"),
        "Flood Zone 3": ("high risk", "high")
    }
    desc, risk_level = risk_dict[flood_zone]
    paragraph_base_dict = {
        "Flood Zone 1": ("According to the GOV.uk flood maps for planning, the site is located in Flood Zone 1, meaning it has less than a 0.1% (1-in-1,000-year) annual probability of river or sea flooding.\n\r"
        "The Environment Agency (EA) flood map for planning is provided in Appendix C.\n\r"
        "Therefore, the risk of the development flooding from fluvial and tidal sources is considered low, and standard planning precautions are sufficient.\n\r"),

        "Flood Zone 2": ("According to the GOV.uk flood maps for planning, the site is located in Flood Zone 2, meaning it has between a 0.1% and 1% (1-in-1,000 to 1-in-100-year) annual probability of river or sea flooding.\n\r"
        "The Environment Agency (EA) flood map for planning is provided in Appendix C.\n\r"
        "Therefore, the risk of the development flooding from fluvial and tidal sources is considered moderate, and appropriate mitigation measures should be considered during the planning and design stages.\n\r"),

        "Flood Zone 3":( "According to the GOV.uk flood maps for planning, the site is located in Flood Zone 3, meaning it has a 1% or greater (1-in-100-year or higher) annual probability of river or sea flooding.\n\r"
        "The Environment Agency (EA) flood map for planning is provided in Appendix C.\n\r"
        "Therefore, the risk of the development flooding from fluvial and tidal sources is considered high, and careful consideration of flood risk management and mitigation measures is strongly recommended.\n\r")

    }
    analysis_text = paragraph_base_dict[flood_zone]

    # ---------- 2. Surface Water Flood Analysis ----------
    easting, northing = postcode_to_easting_northing(postcode)
    if easting is None:
        raise ValueError("无法获取邮编坐标")
    browser = create_browser()
    browser.set_window_position(-2000, 0)
    browser.get(f'https://check-long-term-flood-risk.service.gov.uk/map?easting={easting}&northing={northing}&map=SurfaceWater')
    browser.set_window_size(1920,1200)
    time.sleep(2)
    browser.find_element(By.XPATH, '/html/body/div[1]/form/div/div[1]/div[2]/button[1]').click()
    browser.find_element(By.XPATH, '/html/body/div[1]/form/div/div[2]/div[2]/a').click()
    browser.find_element(By.XPATH, '//*[@id="sw-extent-radio-cc"]').click()
    for _ in range(2):
        browser.find_element(By.XPATH, '//*[@id="map"]/div[1]/div[3]/div[1]/div[4]/div/calcite-button[1]').click()
        time.sleep(1)
    browser.find_element(By.XPATH, '//*[@id="selected-address-checkbox"]').click()
    time.sleep(2)
    png=browser.get_screenshot_as_png()
    browser.quit()
    im = Image.open(io.BytesIO(png))
    x0, y0 = 600, 400
    x1, y1 = im.size[0]-950, im.size[1]-350
    surface_crop = im.crop((x0, y0, x1, y1))


    # 绘制多边形
    surface_image_pil, surface_points=draw_polygon_on_image(surface_crop)

    # ---------- Surface Water DEPTH map ----------
    browser = create_browser()
    browser.set_window_position(-2000, 0)
    browser.get(
        f'https://check-long-term-flood-risk.service.gov.uk/map?easting={easting}&northing={northing}&map=SurfaceWater/depth'
    )
    browser.set_window_size(1920, 1200)
    time.sleep(2)

    browser.find_element(By.XPATH, '/html/body/div[1]/form/div/div[1]/div[2]/button[1]').click()
    browser.find_element(By.XPATH, '/html/body/div[1]/form/div/div[2]/div[2]/a').click()
    browser.find_element(By.XPATH, '//*[@id="sw-extent-radio-cc"]').click()

    for _ in range(2):
        browser.find_element(
            By.XPATH,
            '//*[@id="map"]/div[1]/div[3]/div[1]/div[4]/div/calcite-button[1]'
        ).click()
        time.sleep(1)

    browser.find_element(By.XPATH, '//*[@id="selected-address-checkbox"]').click()


    browser.find_element(By.XPATH, '//*[@id="depth-radio-cc"]').click()

    time.sleep(2)

    png=browser.get_screenshot_as_png()
    browser.quit()

    im = Image.open(io.BytesIO(png))

    # 裁剪 depth 地图

    depth_crop = im.crop((x0, y0, x1, y1))

    surface_depth_img_pil = depth_crop  # 直接用


    # 分析颜色
    surface_rgb = np.array(surface_image_pil)


    low_color = np.array([196,225,255])
    med_color = np.array([154,160,222])
    high_color = np.array([85,92,157])
    if color_in_region(surface_rgb, surface_points, high_color):
        surface_risk = "High chance"
    elif color_in_region(surface_rgb, surface_points, med_color):
        surface_risk = "Medium chance"
    else:
        surface_risk = "Low chance"

    fixed_intro = ("Surface water flooding occurs when intense rainfall is unable to soak into the ground or enter drainage systems, because of blockages or breakages in water pipes or where the drainage capacity has been exceeded. The extent of surface water flooding will depend upon the rainfall event, the degree of saturation of the soil, the permeability of soils and the topography of the site.\n\r")

    # 根据不同风险 level 设置 annual probability
    if risk_level == "High":
        annual_probability = "more than 3.3%"
        assessed_risk = "high"
    elif risk_level == "Medium":
        annual_probability = "between 1% and 3.3%"
        assessed_risk = "medium"
    else:  # Low
        annual_probability = "between 0.1% and 1%"
        assessed_risk = "low"

    surface_text = f"{fixed_intro}A review of the GOV.uk surface water flood risk maps indicates that the site is at ‘{assessed_risk}’ risk of surface water flooding between 2040 and 2060, with an annual probability of {annual_probability}.\n\rBased on this assessment, the risk of flooding from overland surface water flow is considered {assessed_risk}.\n\r"


    # ---------- 3. Reservoirs Flood Analysis ----------
    easting, northing = postcode_to_easting_northing(postcode)
    if easting is None:
        raise ValueError("无法获取邮编坐标")

    browser = create_browser()
    browser.set_window_position(-2000, 0)
    browser.get(
        f'https://check-long-term-flood-risk.service.gov.uk/map?easting={easting}&northing={northing}&map=Reservoirs')
    browser.set_window_size(1920,1200)
    time.sleep(2)

    # 关闭 cookie banner 和初始弹窗
    browser.find_element(By.XPATH, '/html/body/div[1]/form/div/div[1]/div[2]/button[1]').click()
    browser.find_element(By.XPATH, '/html/body/div[1]/form/div/div[2]/div[2]/a').click()

    # 点击范围按钮
    browser.find_element(By.XPATH, '//*[@id="selected-address-checkbox"]').click()

    # 放大地图两次
    for _ in range(2):
        browser.find_element(By.XPATH, '//*[@id="map"]/div[1]/div[3]/div[1]/div[4]/div/calcite-button[1]').click()
        time.sleep(1)

    # 截图
    png = browser.get_screenshot_as_png()
    browser.quit()

    # 裁剪地图
    im = Image.open(io.BytesIO(png))
    x0, y0 = 600, 400
    x1, y1 = im.size[0] - 950, im.size[1] - 350
    reservoir_crop = im.crop((x0, y0, x1, y1))
    reservoir_image_pil, reservoir_points = draw_polygon_on_image(reservoir_crop)


    # 分析颜色
    reservoir_rgb = np.array(reservoir_image_pil)

    medium_color = np.array([85, 92, 157])  # #555C9D
    low_color = np.array([196, 225, 255])  # #C4E1FF

    if color_in_region(reservoir_rgb, reservoir_points, medium_color):
        reservoir_risk = "medium"
    else:
        reservoir_risk = "low"

    # 固定介绍段落
    fixed_intro = (
        "Reservoirs are artificially created lakes that are usually formed by building a dam across a river. If one of the dams failed then water could escape from the reservoir, resulting in land or property being flooded.\n\r"
    )

    # 分析段落
    flood_text = (
        f"The EA has mapped areas which could be subject to flooding in the event of reservoir failure. "
        f"A review of the reservoir flood risk map indicated that the site is{' ' if reservoir_risk == 'medium' else ' not '}located within a reservoir Flood Risk Zone "
        f"(an area expects to flood if a local reservoir were to fail or be breached).\n\r"
        f"Following a review of the relevant information, the risk of flooding from reservoirs failure is considered to be {reservoir_risk}.\n\r"
    )

    reservoir_analysis_text = f"{fixed_intro}{flood_text}"


# ---------- 4. 生成报告 ----------
    doc_template = "FRA.docx"
    doc = Document(doc_template)

    info_dict = {
        "<!--PROJECT_NAME-->": project_name,
        "<!--CLIENT_NAME-->": client_name,
        "<!--PROJECT_LOCATION-->": project_location,
        "<!--PROJECT_NUMBER-->": project_number
    }
    for placeholder, text in info_dict.items():
        replace_placeholder_all(doc, placeholder, text)


    # 插入 Fluvial/Tidal 分析
    insert_text_and_image_below(doc, "<!-- INSERT_FLOOD_ANALYSIS_HERE -->",
                               text=analysis_text, image=fluvial_img_pil)

    # 插入 Surface Water 分析文字和图片
    surface_pic_paragraph = insert_text_and_image_below(
        doc,
        "<!-- INSERT_SURFACE_WATER_ANALYSIS_HERE -->",
        text=surface_text,
        image=surface_image_pil
    )

    # 插入 Surface Water DEPTH 图，紧跟上一张图片
    insert_text_and_image_below(
        doc,
        placeholder=None,
        text=None,
        image=surface_depth_img_pil,
        # after_paragraph 参数不再需要，默认紧跟上一段返回的段落
    )

    # 插入 Reservoirs 分析
    insert_text_and_image_below(doc, "<!-- INSERT_RESERVOIRS_ANALYSIS_HERE -->",
                               text=reservoir_analysis_text, image=reservoir_image_pil)


    output_file = f"Flood_Analysis_{postcode}.docx"
    doc.save(output_file)
    print(f"报告已生成: {output_file}")



